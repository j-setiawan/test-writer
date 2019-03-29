from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from collections import OrderedDict
from math import ceil

from model.test import *


def set_answer_space(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    tag = 'w:{}'.format('bottom')

    element = tc_borders.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        tc_borders.append(element)

    element.set(qn('w:color'), '#000000')
    element.set(qn('w:val'), 'single')


def add_blank_spacer(document):
    document.add_paragraph(style='EmptyTableRowStyle')


def add_paragraph_spacer(document):
    document.add_paragraph(style='ParagraphSpacerStyle')


def string_to_width(line):
    return Pt(len(line) * 6)  # For inches: * 0.085


def write_roman(num):
    roman = OrderedDict()
    roman[1000] = "M"
    roman[900] = "CM"
    roman[500] = "D"
    roman[400] = "CD"
    roman[100] = "C"
    roman[90] = "XC"
    roman[50] = "L"
    roman[40] = "XL"
    roman[10] = "X"
    roman[9] = "IX"
    roman[5] = "V"
    roman[4] = "IV"
    roman[1] = "I"

    def roman_num(num):
        for r in roman.keys():
            x, y = divmod(num, r)
            yield roman[r] * x
            num -= (r * x)
            if num <= 0:
                break

    return "".join([a for a in roman_num(num)])


test = Test('Test', [])

document = Document()
document.styles.add_style('EmptyTableRowStyle', WD_STYLE_TYPE.PARAGRAPH).font.size = Pt(1)
document.styles.add_style('ParagraphSpacerStyle', WD_STYLE_TYPE.PARAGRAPH).font.size = Pt(8)

document.add_heading(test.title, 0)

test_question_num = 1
sub_q_num = 0

total_points = 0

for q in test.questions:
    if q.sub:
        test_question_num -= 1
        sub_q_num += 1

    if not q.sub and sub_q_num > 0:
        sub_q_num = 0

    if q.points == 0:
        points_string = ''
    elif q.get_num_questions() == 1:
        points_string = '({})'.format(q.points)
    else:
        points_string = '({}x{})'.format(q.points, q.get_num_questions())

    document.add_paragraph().add_run('{}. {} {}'.format(
        write_roman(test_question_num) if not q.sub else write_roman(test_question_num) + '-{}'.format(sub_q_num),
        q.question, points_string)).bold = True

    total_points += q.points * q.get_num_questions()

    if isinstance(q, ShortQuestion):
        num_choices = ceil(len(q.choices) / 2)
        table = document.add_table(rows=num_choices, cols=2)

        choice_num = 1
        for choice in q.choices:
            row_num = ceil(choice_num / 2) - 1
            cell_num = (choice_num - 1) % 2

            table.rows[row_num].cells[cell_num].text = '{}) {}:'.format(choice_num, choice)
            choice_num += 1
    elif isinstance(q, MediumQuestion):
        choice_num = 1
        for choice in q.choices:
            document.add_table(rows=1, cols=1).rows[0].cells[0].paragraphs[0].style = 'EmptyTableRowStyle'

            if not isinstance(choice, list):
                choice = [choice]

            choice_line_num = 0
            for c in choice:
                table = document.add_table(rows=1, cols=2)
                add_blank_spacer(document)
                q_cell = table.rows[0].cells[0]
                q_cell.text = '{} {}'.format('{})'.format(choice_num) if choice_line_num == 0 else '   ', c)
                q_cell.width = string_to_width(c)
                set_answer_space(table.rows[0].cells[1])
                choice_line_num += 1

            choice_num += 1
    elif isinstance(q, LongQuestion):
        choice_num = 1

        for choice in q.choices:
            document.add_paragraph('{}) {}'.format(choice_num, choice))

            for i in range(q.num_answer_lines):
                set_answer_space(document.add_table(rows=1, cols=1).rows[0].cells[0])

            if q.translation:
                translation_table = document.add_table(rows=1, cols=2)
                q_cell = translation_table.rows[0].cells[0]
                q_cell.text = '\nTranslation:'
                q_cell.width = Inches(1)
                set_answer_space(translation_table.rows[0].cells[1])

            add_paragraph_spacer(document)
            choice_num += 1
    elif isinstance(q, FreeFormQuestion):
        choice_num = 1
        for choice in q.choices:
            if not isinstance(choice, list):
                choice = [choice]

            choice_line_num = 0
            for c in choice:
                document.add_paragraph(
                    '{} {}'.format('{})'.format(choice_num) if q.numbered and choice_line_num == 0 else '    ', c))
                choice_line_num += 1

            choice_num += 1
    elif isinstance(q, EssayQuestion):
        answer_table = document.add_table(rows=q.answer_lines, cols=1)

        for x in range(q.answer_lines):
            cell = answer_table.rows[x].cells[0]

            if x > 0:
                cell.text = '\n'

            set_answer_space(cell)
    elif isinstance(q, PictureQuestion):
        table = document.add_table(rows=len(q.choices) * 2, cols=2)
        table.cell(0, 0).merge(table.cell((len(q.choices) * 2) - 1, 0))
        table.rows[0].cells[0].paragraphs[0].add_run().add_picture(q.picture)

        row = 0
        choice_num = 1
        for choice in q.choices:
            table.rows[row].cells[1].text = '{}) {}'.format(choice_num, choice)
            set_answer_space(table.rows[row + 1].cells[1])
            row += 2
            choice_num += 1
    elif isinstance(q, TranslationQuestion):
        document.add_paragraph(q.passage)

        blank_num = 0
        final_question = ''
        split_question = q.fill_in_the_blanks.split('()')
        for part in split_question[:-1]:
            final_question += '{}({}{})'.format(part, chr(ord('①') + blank_num), ' ' * 30)

            blank_num += 1

        final_question += split_question[-1]

        document.add_paragraph(final_question).paragraph_format.line_spacing = WD_LINE_SPACING.DOUBLE

    elif isinstance(q, TableQuestion):
        table = document.add_table(rows=len(q.y_headers) + 1, cols=len(q.x_headers) + 1, style='Table Grid')

        y = 1
        for header in q.y_headers:
            p = table.rows[y].cells[0].paragraphs[0]
            p.add_run('\n{}) {}\n'.format(y, header)).bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            y += 1

        x = 1
        for header in q.x_headers:
            p = table.rows[0].cells[x].paragraphs[0]
            p.add_run(header).bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            x += 1
    elif isinstance(q, DialogueQuestion):
        question_num = 1
        for dialogue in q.dialogues:
            line_num = 0

            for line in dialogue.lines:
                if line_num == 0:
                    document.add_paragraph('{}\t{}:  {}'.format(question_num, 'A', line))
                else:
                    document.add_paragraph('\t{}:  {}'.format('A' if line_num % 2 == 0 else 'B', line))
                line_num += 1

            document.add_paragraph(dialogue.question)
            set_answer_space(document.add_table(rows=1, cols=1).rows[0].cells[0])
            add_paragraph_spacer(document)

            question_num += 1
    else:
        print('ERROR! Unknown question type {}'.format(q))

    add_paragraph_spacer(document)
    test_question_num += 1

document.add_paragraph('    /{}'.format(total_points))
document.save('test.docx')
